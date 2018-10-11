# Global imports
import pandas as pd
import os
from docx import Document
from docx.shared import Inches

# Local import
from settings import facile_test_path

__maintainer__ = 'Pierre Gouedard'


class Documents(object):

    def __init__(self):
        self.path = os.path.join(facile_test_path)

    def save_document(self, document, dirname, type_, name=""):
        if not os.path.exists(os.path.join(self.path, dirname)):
            os.makedirs(os.path.join(self.path, dirname))

        if type_ == 'word':
            document.save(os.path.join(*[self.path, dirname, name]))

        elif type_ == 'excel':
             document.save()

    def Build_word_example(self):

        document = Document()
        document.add_heading('Example', 0)

        p = document.add_paragraph('A plain paragraph having some ')
        p.add_run('bold').bold = True
        p.add_run(' and some ')
        p.add_run('italic.').italic = True

        document.add_heading('Heading, level 1', level=1)
        document.add_paragraph('Intense quote', style='Intense Quote')

        document.add_paragraph(
            'first item in unordered list', style='List Bullet'
        )
        document.add_paragraph(
            'first item in ordered list', style='List Number'
        )

        records = (
            (3, '101', 'Spam'),
            (7, '422', 'Eggs'),
            (4, '631', 'Spam, spam, eggs, and spam')
        )

        table = document.add_table(rows=1, cols=3)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Qty'
        hdr_cells[1].text = 'Id'
        hdr_cells[2].text = 'Desc'
        for qty, id, desc in records:
            row_cells = table.add_row().cells
            row_cells[0].text = str(qty)
            row_cells[1].text = id
            row_cells[2].text = desc

        document.add_page_break()

        self.save_document(document, 'example', 'word', 'example_word.docx')

    def Build_excel_example(self):

        # get a csv
        df1 = pd.read_csv(os.path.join(self.path, 'devis.csv'))
        df2 = pd.read_csv(os.path.join(self.path, 'facture.csv'))

        writer = pd.ExcelWriter(os.path.join(self.path, 'example', 'example_excel.xlsx'))

        df1.to_excel(writer, 'Devis')
        df2.to_excel(writer, 'Facture')

        self.save_document(writer, 'example', 'excel')


if __name__ == '__main__':
    test = Documents()
    test.Build_word_example()
    test.Build_excel_example()
