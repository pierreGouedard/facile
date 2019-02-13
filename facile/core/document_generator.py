# Global import
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from pandas import Timestamp
# Local import


class Documents(object):

    def __init__(self, path, driver):
        self.path = path
        self.driver = driver
        self.document = None

    def save_document(self, name):
        if self.document is not None:
            self.document.save(self.driver.join(self.path, name))
        else:
            raise ValueError('Document is None')


class WordDocument(Documents):
    def __init__(self, path, driver, document_settings):
        Documents.__init__(self, path, driver)
        
        # Create document 
        self.document = Document()
        
        # Set document global parameter
        self.field_size = document_settings.get('field_size', 40)

        # Set margin
        sec = self.document.sections[0]
        sec.top_margin = Inches(document_settings.get('top_margin', 1.))
        sec.left_margin = Inches(document_settings.get('left_margin', 0.5))
        sec.right_margin = Inches(document_settings.get('right_margin', 0.5))

        # Set font type
        style = self.document.styles['Normal']
        font = style.font
        font.name = document_settings.get('font_name', 'DejaVu Sans Mono')
        font.size = Pt(int(document_settings.get('font_size', 8)))

    def save_document(self, name):

        # Add footer to document
        self.add_footer('Edition du document: {}'.format(Timestamp.now().date()))
        Documents.save_document(self, name)

    def add_title(self, title, font_size=12, text_align='center', color='000000', left_indent=0.,
                  space_before=0.12, space_after=0.12):

        h = self.document.add_heading(title, 1)
        h.paragraph_format.left_indent = Inches(left_indent)

        if text_align == 'center':
            h.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

        elif text_align == 'left':
            h.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

        if text_align == 'right':
            h.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        h.style.font.bold = True
        h.style.font.color.rgb = RGBColor.from_string(color)
        h.style.font.size = Pt(font_size)
        h.paragraph_format.space_before = Inches(space_before)
        h.paragraph_format.space_after = Inches(space_after)

    def add_table(self, df, index_column=-1, left_indent=0.15):

        # Load values from dataframe
        l_col_names = df.columns
        l_values = [[row[c] for c in l_col_names] for _, row in df.iterrows()]

        table = self.document.add_table(rows=1, cols=len(l_col_names))
        table.style.paragraph_format.left_indent = Inches(left_indent)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Build header row
        row = table.rows[0]
        l_cells = row.cells
        for cell, name in zip(l_cells, l_col_names):
            cell.text = name
            cell.paragraphs[0].runs[0].bold = True
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        # Build body rows
        for l_row_values in l_values:
            l_cells = table.add_row().cells
            for i, (cell, name) in enumerate(zip(l_cells, l_row_values)):
                cell.text = str(name)
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                cell.vertical_alignment = WD_ALIGN_PARAGRAPH.CENTER
                if i == index_column:
                    cell.paragraphs[0].runs[0].bold = True

    def add_field(self, title, value, left_indent=0., space_before=0.06, space_after=0.):

        # create and format paragraph
        p = self.document.add_paragraph()
        p.paragraph_format.left_indent = Inches(left_indent)
        p.paragraph_format.space_before = Inches(space_before)
        p.paragraph_format.space_after = Inches(space_after)
        tab = ' ' + " ".join(['.'] * ((self.field_size - len(title)) / 2)) + ' ' * (len(title) % 2)

        # Add info
        p.add_run(title).bold = True
        p.add_run(tab)
        p.add_run(str(value))

    def add_simple_paragraph(self, l_runs, break_run=False, left_indent=0., space_before=0.06, space_after=0.,
                             bold=False, alignment=None):

        p = self.document.add_paragraph()

        if alignment == 'center':
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif alignment == 'right':
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        p.paragraph_format.left_indent = Inches(left_indent)
        p.paragraph_format.space_before = Inches(space_before)
        p.paragraph_format.space_after = Inches(space_after)

        for text in l_runs:
            r = p.add_run(text)
            r.bold = bold

            if break_run:
                r.add_break()

    def add_footer(self, text):
        footer = self.document.sections[0].footer
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        r.add_break()
        p.add_run('CASOE')
