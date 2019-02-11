# Global import
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT

# Local import
from facile.core.table_loader import TableLoader
from facileapp.models.views.base_view import BaseView
from facileapp.models.client import Client
from facileapp.models.chantier import Chantier
from facileapp.models.contact import Contact
from facileapp.models.devis import Devis
from facileapp.models.affaire import Affaire
from facileapp.models.facture import Facture
from facileapp.models.commande import Commande
from facileapp.models.heure import Heure


class FeuilleTravaux(BaseView):
    main_model = Affaire
    l_models = [Affaire,  Devis, Facture, Commande, Heure]
    l_main_index = [f.name for f in Affaire.l_index]

    @staticmethod
    def merge_to_main(df_main, df_sub, l_col_sub, d_fillna=None):

        l_col_main = FeuilleTravaux.l_main_index

        df_sub = df_sub[l_col_main + l_col_sub] \
            .groupby(l_col_main) \
            .sum() \
            .reset_index()

        df_main = df_main.merge(df_sub, on=l_col_main, how='left')

        if d_fillna is None:
            d_fillna = {c: 0.0 for c in l_col_sub}

        df_main = df_main.fillna(d_fillna)

        return df_main

    @staticmethod
    def split_main_indices(df):
        df['affaire_num'] = df.affaire_id.apply(lambda x: x.split('-')[0])
        df['affaire_ind'] = df.affaire_id.apply(lambda x: x.split('-')[1])

        return df

    @staticmethod
    def load_view():
        # Load affaire db
        df = Affaire.load_db()

        # Join devis information
        df_devis = Devis.load_db()
        df_devis = df_devis[[c for c in df_devis.columns if c not in ['creation_date', 'maj_date']]]
        df_devis = df_devis.rename(columns={c: '{}_devis'.format(c) for c in df_devis.columns if c != 'devis_id'})
        df = df.merge(df_devis, on='devis_id', how='left')

        # Load billing and Build affaire index
        df_facture = Facture.load_db()
        df_facture = FeuilleTravaux.split_main_indices(df_facture)

        # For each situation join amount factured
        for i in range(1, 13):
            df_facture_ = df_facture.loc[df_facture['situation'] == i]\
                .rename(columns={'montant_ht': 'montant_situation_{}'.format(i)})
            df = FeuilleTravaux.merge_to_main(df, df_facture_, ['montant_situation_{}'.format(i)])

        # Join command information
        df_commande = Commande.load_db()
        df_commande = FeuilleTravaux.split_main_indices(df_commande)

        df = FeuilleTravaux\
            .merge_to_main(
                df, df_commande.rename(columns={'montant_ht': 'montant_total_commande'}), ['montant_total_commande']
            )

        # Load hours information
        df_heure = Heure.load_db()
        df_heure = FeuilleTravaux.split_main_indices(df_heure)

        # Heures non interimaires
        df_heure_ = df_heure.loc[df_heure['name'] != 'interim']\
            .rename(columns={'heure_prod': 'heure_prod_saisie', 'heure_autre': 'heure_autre_saisie'})
        df = FeuilleTravaux.merge_to_main(df, df_heure_, ['heure_prod_saisie', 'heure_autre_saisie'])

        # Heures interimaires
        df_heure_ = df_heure.loc[df_heure['name'] == 'interim']\
            .rename(columns={'heure_prod': 'heure_prod_saisie_interim', 'heure_autre': 'heure_autre_saisie_interim'})
        df = FeuilleTravaux.merge_to_main(df, df_heure_, ['heure_prod_saisie_interim', 'heure_autre_saisie_interim'])

        return df

    @staticmethod
    def html_table_loading():
        l_index = FeuilleTravaux.main_model.l_index
        l_fields = FeuilleTravaux.main_model.l_fields()

        # Load database
        df = FeuilleTravaux.load_view()

        l_model_cols = [f.name for f in l_index + l_fields]
        l_extra_cols = [c for c in df.columns if c not in l_model_cols]
        table_man = TableLoader(l_index, l_fields)
        df, d_footer, kwargs = table_man.load_full_table(df, l_extra_cols=l_extra_cols)

        return df, d_footer, kwargs

    @staticmethod
    def document_loading(key, index):
        df = FeuilleTravaux.load_view()
        df = df.loc[
            df[[f.name for f in Affaire.l_index]].apply(lambda r: all([r[c] == index[c] for c in r.index]), axis=1)
        ]

        # Load contact
        df_contact = Contact.load_db()

        # Load client
        df_client = Client.load_db()

        # Load chantier
        df_chantier = Chantier.load_db()

        if key == 'ftravaux':

            document = Document()

            sec = document.sections[0]
            sec.top_margin = Inches(1)
            sec.left_margin = Inches(0.5)
            sec.right_margin = Inches(0.5)
            line_size = 40
            style = document.styles['Normal']
            font = style.font
            font.name = 'DejaVu Sans Mono'
            font.size = Pt(8)

            title = 'Feuille de travaux {}'.format('/'.join(df[FeuilleTravaux.l_main_index].values[0]))
            _ = add_title(document, title, font_size=15, text_align='center', color='000000')

            ########### CLIENT
            _ = add_title(document, 'Client', font_size=12, text_align='left', color='000000')
            s_client = df_client.loc[df_client.designation == df['designation_client_devis'].iloc[0]].iloc[0]

            title = 'Designation'
            value = s_client['designation']
            add_field(document, title, value, line_size=line_size, left_indent=0.15)

            title = 'Raison sociale'
            value = s_client['raison_social']
            add_field(document, title, value, line_size=line_size, left_indent=0.15)

            title = 'Adresse'
            value = '{}, {} - {} {}'.format(
                s_client['adresse'], s_client['cs_bp'], s_client['code_postal'], s_client['ville']
            )
            add_field(document, title, value, line_size=line_size, left_indent=0.15)

            title = 'Responsable commande'
            value = df_contact.loc[df_contact.contact_id == df['contact_id_devis'].iloc[0], 'contact'].iloc[0]
            add_field(document, title, value, line_size=line_size, left_indent=0.15)

            ########### DEVIS
            _ = add_title(document, 'Devis', font_size=12, text_align='left', color='000000')

            title = 'Numero'
            value = df['devis_id'].iloc[0]
            add_field(document, title, value, line_size=line_size, left_indent=0.15)

            title = 'Objet'
            value = df['object_devis'].iloc[0]
            add_field(document, title, value, line_size=line_size, left_indent=0.15)

            title = 'Responsable devis'
            value = df['responsable_devis'].iloc[0]
            add_field(document, title, value, line_size=line_size, left_indent=0.15)

            title = 'Montant total du devis'
            value = '{} euros'.format(df['price_devis'].iloc[0])
            add_field(document, title, value, line_size=line_size, left_indent=0.15)

            add_simple_paragraph(
                document, ['Details'], space_before=0.06, space_after=0.06, left_indent=0.15, bold=True
            )

            l_values = [[df['heure_prod_devis'].iloc[0], df['prix_heure_prod_devis'].iloc[0],
                         df['heure_autre_devis'].iloc[0], df['prix_heure_autre_devis'].iloc[0],
                         df['montant_achat_devis'].iloc[0], df['coef_achat_devis'].iloc[0]]]

            df_table = pd.DataFrame(
                l_values, columns=['Heures Prod', 'Prix Heures Prod', 'Heures Autres', 'Prix Heures Autres',
                                   'Montant achat', 'Coef achat']
            )
            add_table(document, df_table, index_column=-1, left_indent=0.15)

            ########### CHANTIER
            _ = add_title(document, 'Chantier', font_size=12, text_align='left', color='000000')
            s_chantier = df_chantier.loc[df_chantier['chantier_id'] == df['chantier_id'].iloc[0]].iloc[0]
            contact_client_ch = df['contact_chantier_client'].iloc[0]
            designation = df_contact.loc[df_contact.contact_id == contact_client_ch, 'contact'].iloc[0]

            coord = '{}, {} {}'.format(
                s_chantier['adresse'], s_chantier['code_postal'], s_chantier['ville']
            )

            title = 'Adresse'
            value = coord
            add_field(document, title, value, line_size=line_size, left_indent=0.15)

            title = 'Responsable interne'
            value = df['contact_chantier_interne'].iloc[0]
            add_field(document, title, value, line_size=line_size, left_indent=0.15)

            title = 'Responsable client'
            value = '{} ({})'.format(contact_client_ch, designation)
            add_field(document, title, value, line_size=line_size, left_indent=0.15)

            ########### SUIVI
            _ = add_title(document, 'Suivi', font_size=12, text_align='left', color='000000')
            heure_prod = df['heure_prod_saisie'].sum() + df['heure_prod_saisie_interim'].sum()
            heure_autre = df['heure_autre_saisie'].sum() + df['heure_autre_saisie_interim'].sum()
            montant_achat = df['montant_total_commande'].sum()

            l_values = [["REALISE", heure_prod, heure_autre, montant_achat],
                        ["ECART DEVIS", heure_prod - df['heure_prod_devis'].sum(),
                         heure_autre - df['heure_autre_devis'].sum(),
                         montant_achat - df['montant_achat_devis'].sum()]]

            df_table = pd.DataFrame(l_values, columns=[' ', 'Heures prod', 'Heures autre', 'Achat'])

            add_table(document, df_table, index_column=0, left_indent=0.15)

            ########### FACTURATION
            _ = add_title(document, 'Facturation', font_size=12, text_align='left', color='000000')
            s_contact = df_contact.loc[df_contact.contact_id == df['contact_chantier_client'].iloc[0]].iloc[0]
            coord = '{}, {} - {} {}'.format(
                s_contact['adresse'], s_contact['cs_bp'], s_contact['code_postal'], s_contact['ville']
            )

            add_simple_paragraph(
                document, [s_contact['designation'], s_contact['contact'], coord], break_run=True, space_before=0.06,
                alignment='center'
            )

            l_values = [['Visa'] + [str('__/__/____')] * 6, ['Encaissement'] + [str('__/__/____')] * 6]
            df_table = pd.DataFrame(
                l_values, columns=['Situation'] + ['{}'.format(i + 1) for i in range(6)]
            )
            add_table(document, df_table, index_column=0, left_indent=0.15)

            l_values = [['Visa'] + [str('__/__/____')] * 6, ['Encaissement'] + [str('__/__/____')] * 6]

            df_table = pd.DataFrame(
                l_values, columns=['Situation'] + ['{}'.format(i + 7) for i in range(6)]
            )
            add_table(document, df_table, index_column=0, left_indent=0.15)

            document.save('test.docx')

    @staticmethod
    def control_loading():
        d_control_data = {}
        df = FeuilleTravaux.load_view()
        df['affaire_id'] = df[['affaire_num', 'affaire_ind']]\
            .apply(lambda r: '{} - {}'.join([r['affaire_num'], r['affaire_ind']]), axis=1)

        d_name = {True: 'Cloture', False: 'En cours'}
        df_state = df[['affaire_id', 'montant_encaisse', 'montant_facture']]
        df_state['state'] = df_state[['montant_encaisse', 'montant_facture']]\
            .apply(lambda r: d_name[r['montant_encaisse'] == r['montant_facture']], axis=1)

        # App 1 repartition statue of Affaires
        df_state = df_state[['affaire_id', 'state']].groupby('state')\
            .count()\
            .reset_index()\
            .rename(columns={'state': 'name', 'affaire_id': 'value'})

        d_control_data['repstate'] = {
            'plot': {'k': 'pie', 'd': df_state, 'o': {'hover': True}},
            'rows': [('title', [{'content': 'title', 'value': 'Repartition des affaire', 'cls': 'text-center'}]),
                     ('figure', [{'content': 'plot'}])],
            'rank': 0
                }

        # App 2 amount of digned affaire by charge d'aff
        df_chardaff = df[['responsable', 'montant_facture']].groupby('responsable')\
            .sum()\
            .reset_index()\
            .rename(columns={'responsable': 'label'})

        d_control_data['affaireresp'] = {
            'plot': {'k': 'bar', 'd': df_chardaff, 'o': {'val_col': 'montant_facture'}},
            'rows': [('title', [{'content': 'title', 'value': "Affaire facture par charge d'affaire", 'cls': 'text-center'}]),
                     ('figure', [{'content': 'plot'}])],
            'rank': 1
        }

        # App 3 amount of digned affaire by cient
        df_client = df[['rs_client_devis', 'montant_facture']].groupby('rs_client_devis') \
            .sum() \
            .reset_index() \
            .rename(columns={'rs_client_devis': 'label'})

        d_control_data['affaireclient'] = {
            'plot': {'k': 'bar', 'd': df_client, 'o': {'val_col': 'montant_facture'}},
            'rows': [('title',
                      [{'content': 'title', 'value': "Affaire facture par client", 'cls': 'text-center'}]),
                     ('figure', [{'content': 'plot'}])],
            'rank': 2
        }

        return d_control_data


def add_title(document, title, font_size=12, text_align='center', color='000000', left_indent=0.):

    h = document.add_heading(title, 1)
    h.paragraph_format.left_indent = Inches(left_indent)

    if text_align == 'center':
        h.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    elif text_align == 'left':
        h.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    if text_align == 'right':
        h.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    h.style.font.bold = True
    h.style.font.color.rgb = RGBColor.from_string(color)
    h.style.font.size = Pt(font_size)
    h.paragraph_format.space_before = Inches(0.12)
    h.paragraph_format.space_after = Inches(0.12)
    return h


def add_table(document, df, index_column=-1, left_indent=0.15):

    # Load values from dataframe
    l_col_names = df.columns
    l_values = [[row[c] for c in l_col_names] for _, row in df.iterrows()]

    table = document.add_table(rows=1, cols=len(l_col_names))
    table.style.paragraph_format.left_indent = Inches(left_indent)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Build header row
    row = table.rows[0]
    l_cells = row.cells
    for cell, name in zip(l_cells, l_col_names):
        cell.text = name
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Build body rows
    for l_row_values in l_values:
        l_cells = table.add_row().cells
        for i, (cell, name) in enumerate(zip(l_cells, l_row_values)):
            cell.text = str(name)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.vertical_alignment = WD_ALIGN_PARAGRAPH.CENTER
            if i == index_column:
                cell.paragraphs[0].runs[0].bold = True


def add_field(document, title, value, line_size, left_indent=0., space_before=0.06, space_after=0.):

    # create and format paragraph
    p = document.add_paragraph()
    p.paragraph_format.left_indent = Inches(left_indent)
    p.paragraph_format.space_before = Inches(space_before)
    p.paragraph_format.space_after = Inches(space_after)
    tab = ' ' + " ".join(['.'] * ((line_size - len(title)) / 2)) + ' ' * (len(title) % 2)

    # Add info
    p.add_run(title).bold = True
    p.add_run(tab)
    p.add_run(str(value))

    return p


def add_simple_paragraph(document, l_runs, break_run=False, left_indent=0., space_before=0.06, space_after=0.,
                         bold=False, alignment=None):

    p = document.add_paragraph()

    if alignment == 'center':
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif alignment == 'right':
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    p.paragraph_format.left_indent = Inches(left_indent)
    p.paragraph_format.space_before = Inches(space_before)
    p.paragraph_format.space_after = Inches(space_after)

    for text in l_runs:
        r = p.add_run(text)
        r.bold = bold

        if break_run:
            r.add_break()

    return p